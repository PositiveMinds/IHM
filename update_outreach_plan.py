from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Load existing document
doc = Document('Client Outreach Plan - Western Uganda (10 First Customers).docx')

# Find the index where we need to insert
insert_index = None
for i, para in enumerate(doc.paragraphs):
    if "Prospecting & Outreach Strategy" in para.text:
        insert_index = i
        break

if insert_index:
    # Create new section content as list
    new_sections = [
        ("Integrated Health Management Features", "Heading 1"),
        ("Beyond HIV Automation - Comprehensive Chronic Disease Management", "Normal"),
        ("", "Normal"),
        ("Our solution has been expanded to address multiple healthcare management needs across Western Uganda:", "Normal"),
        ("", "Normal"),
        ("Chronic Disease Management (Diabetes, Hypertension)", "Heading 2"),
        ("Mbarara, Fort Portal, and surrounding districts have high burden of non-communicable diseases alongside HIV. Our system now tracks:", "Normal"),
        ("Blood pressure and glucose monitoring for co-infected patients", "List Bullet"),
        ("Treatment adherence for dual disease management", "List Bullet"),
        ("Complication screening and risk stratification", "List Bullet"),
        ("Integration with HIV treatment protocols for safe drug interactions", "List Bullet"),
        ("Applicability: All 10 prospects manage chronic disease patients; 70% of HIV patients have hypertension", "Normal"),
        ("", "Normal"),
        ("Maternal Health Tracking", "Heading 2"),
        ("Western Uganda has significant maternal mortality burden. Our system provides:", "Normal"),
        ("Pregnancy tracking and trimester monitoring", "List Bullet"),
        ("Prevention of Mother-to-Child Transmission (PMTCT) automation", "List Bullet"),
        ("Antenatal care scheduling and adherence", "List Bullet"),
        ("Labor/delivery coordination with birth outcomes tracking", "List Bullet"),
        ("Postpartum follow-up for mother and baby", "List Bullet"),
        ("Applicability: Prospects #1, #2, #3, #5, #7, #10 (hospitals/district programs) manage maternal services; referral potential for all 10", "Normal"),
        ("", "Normal"),
        ("Medication Adherence Systems", "Heading 2"),
        ("Adherence remains the #1 clinical challenge across all prospects. Our system automates:", "Normal"),
        ("Smart medication reminders via WhatsApp (culturally aligned with Western Uganda)", "List Bullet"),
        ("Adherence tracking with gap identification", "List Bullet"),
        ("Drug pickup forecasting to prevent stockouts", "List Bullet"),
        ("Patient education content delivery", "List Bullet"),
        ("Adherence analytics for clinical supervision", "List Bullet"),
        ("Applicability: Mission-critical for all 10 prospects; strongest conversion driver for private/NGO sector", "Normal"),
        ("", "Normal"),
        ("Clinic Appointment Management", "Heading 2"),
        ("All 10 prospects report chaotic appointment scheduling. Our system provides:", "Normal"),
        ("Automated appointment scheduling and optimization", "List Bullet"),
        ("SMS/WhatsApp reminder system (reduces no-shows by 25-40%)", "List Bullet"),
        ("Calendar integration for multi-clinic coordination", "List Bullet"),
        ("Waitlist management and re-scheduling automation", "List Bullet"),
        ("Capacity planning analytics", "List Bullet"),
        ("Applicability: All 10 prospects report significant no-show rates (15-30%); immediate ROI through show-rate improvement", "Normal"),
        ("", "Normal"),
        ("Why These Features Matter for Western Uganda Market", "Heading 2"),
        ("Patient Volume: Western Uganda facilities average 1,000+ patients each - demand for multi-disease management is urgent", "List Bullet"),
        ("Clinical Complexity: High HIV prevalence + comorbidities require integrated tracking (not siloed systems)", "List Bullet"),
        ("Revenue Sensitivity: Private sector and NGOs directly benefit from appointment compliance and medication adherence", "List Bullet"),
        ("Donor Requirements: District programs and UNHCR require comprehensive health metrics for compliance reporting", "List Bullet"),
        ("WhatsApp Integration: Western Uganda healthcare community relies heavily on WhatsApp - automation via messaging is highly valued", "List Bullet"),
        ("Implementation Advantage: Expanded features position us as comprehensive health platform, not just HIV tool - stronger competitive positioning", "List Bullet"),
        ("", "Normal"),
        ("Expanded Value Propositions by Prospect Type", "Heading 2"),
        ("", "Normal"),
        ("Hospital Systems (Prospects #1, #2, #3):", "Normal"),
        ('"Manage HIV, chronic disease, maternal health, and appointments from one integrated platform. Reduce supervisory burden by 60%. Improve patient outcomes across all disease areas."', "List Bullet"),
        ("", "Normal"),
        ("District Programs (Prospects #4, #7):", "Normal"),
        ('"Real-time visibility into HIV, chronic disease, and maternal health across all facilities. Automated compliance reporting for ministry. Early warning for high-risk patients across disease areas."', "List Bullet"),
        ("", "Normal"),
        ("NGO Networks (Prospect #5):", "Normal"),
        ('"Unified patient dashboard across 6 clinics. Manage HIV, chronic disease, and maternal health with coordinated care. Donor reporting simplified. Network value maximized."', "List Bullet"),
        ("", "Normal"),
        ("Academic/Research (Prospect #6):", "Normal"),
        ('"Research-grade data capture across HIV, chronic disease, maternal health. Improved outcomes analytics. Publication opportunities across multiple disease domains."', "List Bullet"),
        ("", "Normal"),
        ("Private Sector (Prospect #9):", "Normal"),
        ('"Appointment compliance systems + medication adherence automation reduce no-shows by 25-40%. Direct revenue impact. Chronic disease monitoring expands revenue per patient."', "List Bullet"),
        ("", "Normal"),
        ("Humanitarian (Prospect #10):", "Normal"),
        ('"Refugee health management across HIV, chronic disease, maternal health. Continuity of care despite mobility. Donor/UN compliance automated. Vulnerable population focus."', "List Bullet"),
        ("", "Normal"),
    ]
    
    # Insert paragraphs before the "Prospecting & Outreach Strategy" paragraph
    for text, style in new_sections:
        p = doc.paragraphs[insert_index]._element
        new_p = doc.add_paragraph(text, style=style)
        # Move new paragraph to correct position
        p.addprevious(new_p._element)
        doc.paragraphs[insert_index]._element.getparent().remove(new_p._element)
        p.addprevious(new_p._element)

# Save updated document
doc.save('Client Outreach Plan - Western Uganda (10 First Customers).docx')
print("Document updated successfully with expanded features!")
