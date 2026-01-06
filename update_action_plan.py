from docx import Document

# Update 90-Day Action Plan document
doc = Document('90-Day Action Plan - AI HIV Automation (Western Uganda).docx')

# Update main title
if len(doc.paragraphs) > 0:
    doc.paragraphs[0].text = "Integrated Healthcare Management Platform"
    doc.paragraphs[0].style = 'Heading 1'

# Update subtitle
if len(doc.paragraphs) > 1:
    doc.paragraphs[1].text = "HIV, Chronic Disease, Maternal Health & Appointment Management - 90-Day Action Plan"
    doc.paragraphs[1].style = 'Heading 2'

# Save
doc.save('90-Day Action Plan - AI HIV Automation (Western Uganda).docx')
print("90-Day Action Plan updated!")
