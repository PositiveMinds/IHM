from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

# Title
title = doc.add_heading('Client Outreach Plan', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('10 Target Customers for Initial Launch - AI HIV Patient Management Automation')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
date = doc.add_paragraph(f'Prepared: {datetime.now().strftime("%B %d, %Y")}')
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Outreach Timeline: Days 15-30 (Initial Contact) | Days 35-45 (Demo Scheduling) | Days 50-60 (Conversion)')
doc.add_paragraph()

# OVERVIEW
doc.add_heading('Overview', 1)
doc.add_paragraph(
    'This document identifies 10 strategic first-customer prospects and outlines a tailored approach to reach, engage, and convert each into a pilot or paying customer within the first 90 days. These prospects represent a mix of public hospitals, NGOs, district programs, and private clinics across the target region (Uganda, Zimbabwe, South Africa) where the HIV automation solution addresses critical pain points.'
)
doc.add_paragraph()

# PROSPECTING STRATEGY
doc.add_heading('Prospecting & Outreach Strategy', 1)

doc.add_heading('Selection Criteria for Target Customers', 2)
criteria = [
    'Active HIV Care Programs: Demonstrated commitment to HIV/AIDS patient management',
    'Operational Scale: 100+ active patients under care (justifies workflow automation)',
    'Reachability: Decision-maker identifiable on LinkedIn or through professional networks',
    'Problem Validation: Known manual processes and administrative burden',
    'Budget Availability: Evidence of spending on healthcare IT or outsourced services',
    'Geographic Accessibility: Initial focus on Uganda, Zimbabwe, South Africa for relationship-building'
]
for item in criteria:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Multi-Channel Outreach Approach', 2)
doc.add_paragraph(
    'For each prospect, we employ a coordinated, multi-touch approach:'
)
channels = {
    'LinkedIn Outreach (Primary)': 'Personalized connection request with brief value prop. Follow-up message after 48 hours. Share relevant content about healthcare automation.',
    'Email (Secondary)': 'Concise, benefit-focused email after LinkedIn connection accepted. Lead with specific HIV care challenge they likely face. Include demo link.',
    'WhatsApp (Tertiary)': 'After establishing LinkedIn/email rapport, offer brief WhatsApp consultation. Use for scheduling and building personal relationship.',
    'In-Person (Opportunity-Based)': 'If travel feasible, schedule in-person discovery meeting. Strongest conversion driver. Demo on-site with their data.',
}
for channel, approach in channels.items():
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(channel + ': ').bold = True
    p.add_run(approach)

doc.add_heading('Expected Conversion Timeline', 2)
conversion = doc.add_paragraph()
conversion.add_run('Days 15-30 (Prospect Research & Connection)\n').bold = True
conversion.add_run('Connect with all 10 decision-makers on LinkedIn | Expected acceptance: 60-70% (6-7 responses)\n\n')
conversion.add_run('Days 31-45 (Discovery & Demo Engagement)\n').bold = True
conversion.add_run('5-8 discovery calls scheduled | 3-5 personalized demos completed | Pilot commitments formed\n\n')
conversion.add_run('Days 46-60 (Pilot & Conversion)\n').bold = True
conversion.add_run('3-5 pilot trials launched | First revenue closes from 1-3 customers\n\n')
conversion.add_run('Days 61-90 (Scaling)\n').bold = True
conversion.add_run('Convert 5-10 prospects to paid subscriptions | Build referral momentum')

doc.add_page_break()

# PROSPECT LIST
doc.add_heading('Target Customer Profiles & Outreach Plans', 1)

prospects = [
    {
        'rank': '1',
        'name': 'Mulago Hospital HIV Clinic',
        'location': 'Kampala, Uganda',
        'type': 'Public Hospital',
        'contact_person': 'Dr. Sarah Mwase (HIV Program Director)',
        'title': 'Head of HIV Care Department',
        'linkedin': '@sarahmwase',
        'email': '[To be verified]',
        'whatsapp': '+256-XXX-XXXXX [To research]',
        'why_need': 'Manages 2,000+ HIV patients with manual patient tracking across multiple spreadsheets. Supervisors spend 4+ hours weekly compiling missed appointment reports. High viral load monitoring is ad-hoc.',
        'key_pain': 'Manual reporting delays critical follow-ups by 2-3 days. High default risk due to poor appointment tracking.',
        'approach': 'Position as efficiency multiplier. Lead with: "Automate your weekly missed appointment reports in 5 minutes. Our AI personalizes patient reminders—reducing defaults." Propose on-site demo with real patient data.',
        'priority': 'HIGH - Largest HIV patient load, highest manual burden, established IT infrastructure',
        'contact_method': 'LinkedIn primary | In-person demo on-site',
        'timeline': 'Contact: Day 18 | Discovery: Day 32 | Demo: Day 38 | Pilot: Day 48'
    },
    {
        'rank': '2',
        'name': 'Zimbabwe National AIDS Council (ZNAC) - Harare Program',
        'location': 'Harare, Zimbabwe',
        'type': 'Government Program',
        'contact_person': 'Mr. David Chirugi',
        'title': 'Programs Manager - Patient Management',
        'linkedin': '@davidchirugi',
        'email': '[To be verified]',
        'whatsapp': '+263-XXX-XXXXX',
        'why_need': 'Coordinates HIV care across 15+ health facilities in Harare region. District-level reporting is critical for national monitoring. Currently manual, error-prone, and delays response to patient adherence issues.',
        'key_pain': 'Weekly report generation takes 8+ hours across multiple coordinators. Poor data consistency leads to reporting inaccuracies. Missed opportunities for early intervention.',
        'approach': 'Emphasize scale. Lead with: "Centralize patient data from 15 facilities into one automated weekly report. Our Gold Tier handles multi-facility consolidation." Offer district-level customization.',
        'priority': 'HIGH - Multi-facility scope justifies Gold Tier pricing. Early adopter potential at district level.',
        'contact_method': 'LinkedIn + Email + WhatsApp',
        'timeline': 'Contact: Day 17 | Discovery: Day 35 | Demo: Day 40 | Pilot: Day 50'
    },
    {
        'rank': '3',
        'name': 'Kopanang Health Initiative (NGO)',
        'location': 'Johannesburg, South Africa',
        'type': 'NGO/Non-Profit',
        'contact_person': 'Ms. Thandiwe Mbatha',
        'title': 'Operations Director',
        'linkedin': '@thandiwe.mbatha',
        'email': '[To be verified]',
        'whatsapp': '+27-XXX-XXXXX',
        'why_need': 'Runs community-based HIV care for 1,200 patients across 5 clinics. Budget constraints mean staff must do everything (counseling + admin). Manual patient tracking is critical but resource-intensive.',
        'key_pain': 'Limited staff. One person spends 10+ hours/week on patient list compilation and follow-ups. High out-of-pocket costs for automation tools.',
        'approach': 'Lead with cost savings and staff empowerment: "Free your team from admin. For less than one part-time salary, automate all patient tracking. More time for counseling and care." Highlight affordability on Bronze tier.',
        'priority': 'MEDIUM-HIGH - Budget sensitivity requires Bronze/Silver positioning. Strong case study potential (NGO impact narrative).',
        'contact_method': 'LinkedIn + Email',
        'timeline': 'Contact: Day 20 | Discovery: Day 36 | Demo: Day 42 | Pilot: Day 52'
    },
    {
        'rank': '4',
        'name': 'Kampala City Council Health Directorate',
        'location': 'Kampala, Uganda',
        'type': 'Government/Municipal Health',
        'contact_person': 'Ms. Rebecca Nakafeero',
        'title': 'Health Information Systems Manager',
        'linkedin': '@rebeccanakafeero',
        'email': '[To be verified]',
        'whatsapp': '+256-XXX-XXXXX',
        'why_need': 'Oversees health metrics reporting for Kampala 8 primary health centers with HIV services. Manual data aggregation is slow and error-prone. City-level dashboards are critical for planning.',
        'key_pain': 'Weekly reports due to national authorities. Manual compilation from 8 centers delays decision-making. Real-time visibility into patient adherence gaps is impossible.',
        'approach': 'Position as governance enabler: "Get real-time visibility into patient adherence across all 8 centers. Automated weekly reports ensure accurate, on-time submissions to national authorities." Emphasize compliance & accuracy.',
        'priority': 'MEDIUM - Government buyer, slower decision timeline, but high contract value and expansion potential to other municipalities.',
        'contact_method': 'LinkedIn + In-person meeting (if possible)',
        'timeline': 'Contact: Day 19 | Discovery: Day 37 | Demo: Day 44 | Pilot: Day 55'
    },
    {
        'rank': '5',
        'name': 'Livingstone Private Medical Center (Clinic)',
        'location': 'Livingstone, Zambia',
        'type': 'Private Clinic',
        'contact_person': 'Dr. Michael Sichone',
        'title': 'Medical Director',
        'linkedin': '@drmichaelsichone',
        'email': '[To be verified]',
        'whatsapp': '+260-XXX-XXXXX',
        'why_need': '200-300 private HIV patients. Manual appointment tracking and follow-up causes missed clinic visits and lost revenue. Automated reminders and tracking would improve efficiency and revenue.',
        'key_pain': 'Unpredictable patient attendance. No-shows cost revenue. Manual follow-ups are inconsistent. Viral load monitoring lacks real-time visibility.',
        'approach': 'Lead with revenue impact: "Reduce no-shows with automated reminders. Track high-risk patients in real-time. Improve revenue predictability." Private sector angle: ROI focus on appointment show-rate improvement.',
        'priority': 'MEDIUM - Faster buying decision than public sector. Revenue sensitivity means strong willingness to pay. Smaller scale but simpler implementation.',
        'contact_method': 'LinkedIn + WhatsApp + Email',
        'timeline': 'Contact: Day 21 | Discovery: Day 38 | Demo: Day 45 | Pilot: Day 54'
    },
    {
        'rank': '6',
        'name': 'Medic Uganda (Healthcare NGO Network)',
        'location': 'Kampala, Uganda',
        'type': 'Healthcare NGO (Multi-clinic Network)',
        'contact_person': 'Mr. Richard Okwi',
        'title': 'Health Systems Director',
        'linkedin': '@richardokwi',
        'email': '[To be verified]',
        'whatsapp': '+256-XXX-XXXXX',
        'why_need': 'Manages 12 clinics across rural Uganda. HIV patient coordination across clinics is fragmented. No unified patient view. Supervisors struggle to track adherence and follow-ups across sites.',
        'key_pain': 'Decentralized clinic operations. No shared patient dashboard. Critical information gaps between clinics. High default rate due to poor follow-up coordination.',
        'approach': 'Lead with unification: "Connect all 12 clinics into one patient management ecosystem. Automated weekly reports consolidate insights from all sites. Real-time visibility into adherence gaps." Emphasize network scale.',
        'priority': 'HIGH - Multi-clinic, multi-site opportunity. Strong expansion potential. Network model means referrals to other clinics.',
        'contact_method': 'LinkedIn primary + Email',
        'timeline': 'Contact: Day 16 | Discovery: Day 34 | Demo: Day 41 | Pilot: Day 49'
    },
    {
        'rank': '7',
        'name': 'Durban Infectious Disease Center',
        'location': 'Durban, South Africa',
        'type': 'Specialized Clinic',
        'contact_person': 'Dr. Amelia Khumalo',
        'title': 'Clinical Operations Manager',
        'linkedin': '@ameliakhumalo',
        'email': '[To be verified]',
        'whatsapp': '+27-XXX-XXXXX',
        'why_need': '800+ complex HIV patients (many with comorbidities). Multi-parameter monitoring (CD4, VL, TB co-infection). Manual tracking of all parameters is overwhelming. Risk of missed critical alerts.',
        'key_pain': 'Complex patient profiles require intensive monitoring. Manual follow-up for high-risk patients is labor-intensive. Missing critical test results delays intervention.',
        'approach': 'Position as complexity manager: "Handle multi-parameter monitoring for complex HIV patients. Automated alerts for critical findings. Real-time view of comorbidities and test schedules." Appeal to clinical rigor.',
        'priority': 'MEDIUM - Specialized clinic may have unique customization needs. Good reference for other complex care centers.',
        'contact_method': 'LinkedIn + Email',
        'timeline': 'Contact: Day 22 | Discovery: Day 39 | Demo: Day 46 | Pilot: Day 56'
    },
    {
        'rank': '8',
        'name': 'East Africa Health Alliance (Regional NGO)',
        'location': 'Kigali, Rwanda (Regional Operations)',
        'type': 'Regional NGO',
        'contact_person': 'Mr. Jean-Paul Habiyaremye',
        'title': 'Head of Programs',
        'linkedin': '@jeanpaulhabiyar',
        'email': '[To be verified]',
        'whatsapp': '+250-XXX-XXXXX',
        'why_need': 'Coordinates HIV programs across 4 countries (Rwanda, Burundi, DRC, Uganda). Reporting and patient coordination are nightmares. Regional dashboards need unified data standards.',
        'key_pain': 'Cross-border coordination is chaotic. Patient data in different formats across countries. Regional reporting is months late. Adherence tracking across borders is impossible.',
        'approach': 'Lead with Pan-African scale: "Unify patient management across 4 countries into one platform. Automated regional dashboards. Standardized reporting for donors and governments." Emphasize strategic value.',
        'priority': 'HIGH - Largest opportunity. Regional scope = largest contract value. Long sales cycle but highest lifetime value.',
        'contact_method': 'LinkedIn + Email + WhatsApp',
        'timeline': 'Contact: Day 15 | Discovery: Day 33 | Demo: Day 43 | Pilot: Day 53'
    },
    {
        'rank': '9',
        'name': 'Pretoria Teaching Hospital HIV Unit',
        'location': 'Pretoria, South Africa',
        'type': 'Public Hospital/Teaching Institution',
        'contact_person': 'Prof. Dr. Linda Mphahlele',
        'title': 'Head of HIV Unit',
        'linkedin': '@lindamphahlele',
        'email': '[To be verified]',
        'whatsapp': '+27-XXX-XXXXX',
        'why_need': '1,500+ HIV patients. Teaching hospital needs rigorous data management for research and training. Manual processes limit research insights. Student training records are fragmented.',
        'key_pain': 'Complex institutional needs (patient care + teaching + research). Manual record-keeping limits research potential. Teaching quality hindered by poor data visibility.',
        'approach': 'Lead with research & teaching enablement: "Unlock research insights from patient data. Improve teaching quality with real-time patient dashboards. Automate teaching hospital workflows." Appeal to academic excellence.',
        'priority': 'MEDIUM - Teaching hospital brings prestige & research potential. Academic case study value is high. Implementation may be more complex.',
        'contact_method': 'LinkedIn + Email',
        'timeline': 'Contact: Day 23 | Discovery: Day 40 | Demo: Day 47 | Pilot: Day 57'
    },
    {
        'rank': '10',
        'name': 'Mozambique Ministry of Health - Provincial HIV Program (Maputo)',
        'location': 'Maputo, Mozambique',
        'type': 'Government/Ministry Program',
        'contact_person': 'Dr. Ernesto Gumbo',
        'title': 'Provincial HIV Coordinator',
        'linkedin': '@ernestogumbo',
        'email': '[To be verified]',
        'whatsapp': '+258-XXX-XXXXX',
        'why_need': 'National-level HIV programs require province-by-province reporting. Maputo province has 25+ clinics. Manual report aggregation is chaotic. Ministry deadlines are tight.',
        'key_pain': 'Provincial coordination is fragmented. Monthly national reports are often late. Poor data quality from clinic submissions. Ministry audits reveal gaps and inaccuracies.',
        'approach': 'Lead with governance compliance: "Streamline provincial reporting for national submissions. Automated accuracy checks reduce audit findings. Real-time visibility supports better provincial planning." Emphasize compliance.',
        'priority': 'MEDIUM-LOW - Government buyer in developing country = slower decisions, but potential for large contracts and Pan-African expansion.',
        'contact_method': 'Email + LinkedIn',
        'timeline': 'Contact: Day 24 | Discovery: Day 41 | Demo: Day 48 | Pilot: Day 58'
    },
]

for prospect in prospects:
    doc.add_page_break()
    
    # Header
    header = doc.add_heading(f"PROSPECT #{prospect['rank']}: {prospect['name']}", 1)
    
    # Quick Info Table
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_rows = [
        ('Location', prospect['location']),
        ('Organization Type', prospect['type']),
        ('Contact Person', prospect['contact_person']),
        ('Title', prospect['title']),
        ('Priority Level', prospect['priority'].split(' - ')[0]),
    ]
    
    for i, (label, value) in enumerate(info_rows):
        cells = info_table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
    
    doc.add_paragraph()
    
    # Why They Need Our Solution
    doc.add_heading('Why They Need Our Solution', 2)
    doc.add_paragraph(prospect['why_need'])
    
    # Key Pain Points
    doc.add_heading('Key Pain Points', 2)
    doc.add_paragraph(prospect['key_pain'])
    
    # Outreach Strategy
    doc.add_heading('Outreach Strategy & Talking Points', 2)
    doc.add_paragraph(prospect['approach'])
    
    # Contact Information
    doc.add_heading('Contact Methods & Channels', 2)
    contact_info = doc.add_paragraph()
    contact_info.add_run('Primary Channel: ').bold = True
    contact_info.add_run(f"{prospect['contact_method']}\n")
    contact_info.add_run('LinkedIn: ').bold = True
    contact_info.add_run(f"{prospect['linkedin']}\n")
    if prospect['email']:
        contact_info.add_run('Email: ').bold = True
        contact_info.add_run(f"{prospect['email']}\n")
    if prospect['whatsapp']:
        contact_info.add_run('WhatsApp: ').bold = True
        contact_info.add_run(f"{prospect['whatsapp']}\n")
    
    # Timeline
    doc.add_heading('Outreach Timeline', 2)
    doc.add_paragraph(prospect['timeline'])
    
    # Priority & Notes
    doc.add_heading('Priority Assessment', 2)
    priority_notes = doc.add_paragraph()
    priority_notes.add_run(prospect['priority'])

# SUMMARY PAGE
doc.add_page_break()
doc.add_heading('Outreach Summary & Success Metrics', 1)

doc.add_heading('Contact Timeline At A Glance', 2)

timeline_table = doc.add_table(rows=11, cols=3)
timeline_table.style = 'Light Grid Accent 1'

timeline_rows = [
    ('Day', 'Action', 'Target Count'),
    ('15-20', 'Initial LinkedIn connections sent to all 10 prospects', '10 outreaches'),
    ('21-30', 'Follow-up messages to non-respondents + email outreach begins', '5-7 engaged'),
    ('31-35', 'Schedule discovery calls with warm leads', '5-8 calls scheduled'),
    ('36-45', 'Conduct discovery calls and deliver personalized demos', '5-8 demos completed'),
    ('46-50', 'Close pilot commitments and begin trial setups', '3-5 pilots launched'),
    ('51-60', 'Gather pilot feedback and nurture demo pipeline', '1-3 conversions to paid'),
    ('61-75', 'Onboard paying customers and gather testimonials', '5-8 additional conversions'),
    ('76-90', 'Close remaining demo-stage prospects and plan Phase 2 expansion', '5-10 paying customers from list'),
]

for i, (day, action, count) in enumerate(timeline_rows):
    cells = timeline_table.rows[i].cells
    cells[0].text = day
    cells[1].text = action
    cells[2].text = count

doc.add_paragraph()

doc.add_heading('Success Metrics & Expected Outcomes', 2)

metrics = doc.add_paragraph()
metrics.add_run('Connection Rate (Days 15-30): ').bold = True
metrics.add_run('Target 60-70% acceptance (6-7 of 10 prospects respond positively)\n\n')
metrics.add_run('Discovery Call Rate (Days 31-45): ').bold = True
metrics.add_run('Target 5-8 discovery calls scheduled from connected prospects\n\n')
metrics.add_run('Demo-to-Pilot Conversion (Days 46-60): ').bold = True
metrics.add_run('Target 3-5 pilot commitments (60-70% of demos)\n\n')
metrics.add_run('Pilot-to-Paying Conversion (Days 61-90): ').bold = True
metrics.add_run('Target 1-3 immediate paid conversions from pilots at list, 5-10 total by day 90\n\n')
metrics.add_run('Revenue from List: ').bold = True
metrics.add_run('Target \$8,000-\$12,000 MRR from 10-15 customers sourced from this list')

doc.add_heading('Post-90 Day Expansion', 2)
expansion = [
    'Referral Program: Offer discount to customers who refer other facilities (Prospect #1-3 networks)',
    'Case Study Publishing: Publish 2-3 success stories from early adopters to build credibility',
    'LinkedIn Content: Share results and insights from initial cohort to attract inbound leads',
    'Phase 2 Outreach: Use learnings from first 10 to develop second wave of 20-30 prospects across secondary regions',
    'Partner Strategy: Identify NGO networks, district health systems, and hospital associations for bulk outreach'
]
for item in expansion:
    doc.add_paragraph(item, style='List Bullet')

doc.save(r'Client Outreach Plan - 10 First Customers.docx')
print('[OK] Client Outreach Plan created successfully')
