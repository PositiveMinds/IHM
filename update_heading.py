from docx import Document

# Load document
doc = Document('Client Outreach Plan - Western Uganda (10 First Customers).docx')

# Update main title
if len(doc.paragraphs) > 0:
    doc.paragraphs[0].text = "Integrated Healthcare Management Platform - Client Outreach Plan"
    doc.paragraphs[0].style = 'Heading 1'

# Update subtitle
if len(doc.paragraphs) > 1:
    doc.paragraphs[1].text = "HIV, Chronic Disease Management, Maternal Health & Appointment Scheduling"
    doc.paragraphs[1].style = 'Heading 2'

# Update third line
if len(doc.paragraphs) > 2:
    doc.paragraphs[2].text = "10 Target Healthcare Facilities in Western Uganda (Mbarara, Fort Portal, Kabale, Kisoro, Kanungu, Rukungiri, Ntungamo)"

# Save
doc.save('Client Outreach Plan - Western Uganda (10 First Customers).docx')
print("Document heading updated successfully!")
