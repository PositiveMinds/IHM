from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()
# Set default font and spacing
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def justify_paragraph(paragraph):
    """Apply justified alignment and line spacing"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph

# Title
title = doc.add_heading('AI + n8n Automation for HIV Patient Management', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('90-Day Action Plan & Strategic Roadmap')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
date = doc.add_paragraph(f'Prepared: {datetime.now().strftime("%B %d, %Y")}')
date.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# TABLE OF CONTENTS
doc.add_heading('TABLE OF CONTENTS', 1)
toc_items = [
    '1. Executive Summary',
    '2. Market Validation',
    '3. 30-60-90 Day Action Plan',
    '4. Financial Projections',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# SECTION 1: EXECUTIVE SUMMARY
doc.add_heading('1. EXECUTIVE SUMMARY', 1)

doc.add_heading('Business Overview', 2)
p = doc.add_paragraph(
    'We are launching an AI-powered n8n automation solution specifically designed for healthcare facilities and hospitals managing HIV patient care programs. Our solution automates critical administrative workflows that currently consume 2+ hours per week (104-120 hours annually) of healthcare staff time.'
)
justify_paragraph(p)

doc.add_heading('The Problem We Solve', 2)
p = doc.add_paragraph(
    'Healthcare professionals managing HIV patients with comorbidities face a critical administrative burden: manually tracking multiple health conditions (HIV, diabetes, hypertension, TB), monitoring medication adherence across different therapies, managing maternal health for pregnant patients, tracking scattered clinic appointments, and generating weekly reports. This fragmented manual work consumes 15+ hours per week, diverts time from direct patient care, increases missed appointments, reduces medication adherence, and negatively impacts patient outcomes and engagement across all health conditions.'
)
justify_paragraph(p)

doc.add_heading('Unique Value Proposition', 2)
value_props = [
    'Hyper-specialized for HIV patient management + comorbidity support (not generic automation)',
    'AI-powered personalization for patient communications and health engagement',
    'Integrated modules: Chronic disease management, maternal health tracking, medication adherence, appointment management',
    'Saves 12-15 hours monthly per facility (vs 8-10 with HIV-only solution)',
    'Improves patient adherence, reduces missed appointments, and improves clinical outcomes',
    '40%+ gross margins with strong unit economics',
    'Costs 70% less than competing EHR systems',
    'Flexible n8n architecture with no vendor lock-in'
]
for prop in value_props:
    doc.add_paragraph(prop, style='List Bullet')

doc.add_heading('Integrated Platform Modules', 2)
doc.add_paragraph()
doc.add_paragraph('Core HIV Management Module', style='List Bullet').bold = True
p = doc.add_paragraph(
    'Patient cohort tracking, ART adherence monitoring, missed appointment flagging, viral load tracking, treatment interruption alerts'
)
p.paragraph_format.level = 1

doc.add_paragraph('Chronic Disease Management Module', style='List Bullet').bold = True
p = doc.add_paragraph(
    'Diabetes and hypertension patient tracking, medication monitoring, vital signs integration, comorbidity risk alerts, clinic visit scheduling'
)
p.paragraph_format.level = 1

doc.add_paragraph('Maternal Health Tracking Module', style='List Bullet').bold = True
p = doc.add_paragraph(
    'PMTCT (Prevention of Mother-to-Child Transmission) monitoring, antenatal care scheduling, postnatal follow-up tracking, infant prophylaxis adherence, care coordinator alerts'
)
p.paragraph_format.level = 1

doc.add_paragraph('Medication Adherence System', style='List Bullet').bold = True
p = doc.add_paragraph(
    'AI-powered medication reminders via WhatsApp, adherence pattern analysis, early warning for non-adherence, family/caregiver engagement, personalized intervention recommendations'
)
p.paragraph_format.level = 1

doc.add_paragraph('Clinic Appointment Management', style='List Bullet').bold = True
p = doc.add_paragraph(
    'Automated appointment scheduling, multi-day reminder system, no-show prediction, appointment rescheduling workflow, clinic resource optimization'
)
p.paragraph_format.level = 1

doc.add_paragraph()

doc.add_heading('Target Revenue Goal - First 90 Days', 2)
revenue_goals = doc.add_paragraph()
revenue_goals.add_run('Primary Goal: ').bold = True
revenue_goals.add_run('15-20 paying customers at UGX 2,956,300/month (Silver tier) = UGX 44,344,500 - UGX 59,126,000 monthly recurring revenue\n')
revenue_goals.add_run('Profitability Target: ').bold = True
revenue_goals.add_run('16+ customers (break-even at UGX 18,500,000/month fixed costs)\n')
revenue_goals.add_run('Setup Fee Revenue: ').bold = True
revenue_goals.add_run('UGX 7,400,000-18,500,000 per customer (5-10 customers = UGX 37,000,000-185,000,000)\n')
revenue_goals.add_run('Module Upsell Potential: ').bold = True
revenue_goals.add_run('Additional UGX 370,000-740,000/month per facility for each integrated module (Maternal Health, Chronic Disease, etc.)')

doc.add_page_break()

# SECTION 2: MARKET VALIDATION
doc.add_heading('2. MARKET VALIDATION', 1)

doc.add_heading('Evidence That Target Customers Exist', 2)
p = doc.add_paragraph(
    'Target customers are clearly defined and actively seeking solutions:'
)
justify_paragraph(p)
evidence = [
    'Healthcare Subreddits: Posts documenting frustration with manual data entry and administrative burden',
    'LinkedIn & Professional Forums: Consistent demand for healthcare automation and patient management solutions',
    'Freelancer Platforms: Active job postings on Upwork/Fiverr for "healthcare automation" and "patient list management" at $120-$450/month budgets',
    'Industry Networks: WhatsApp groups, HIVSA (HIV South Africa) networks, and healthcare IT forums show strong network effects',
    'Geographic Validation: Initial research in Sub-Saharan Africa (Uganda, Zimbabwe, South Africa) shows critical need in resource-constrained healthcare systems'
]
for item in evidence:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Competitor Analysis', 2)
p = doc.add_paragraph(
    'Current alternatives include:'
)
justify_paragraph(p)
competitors = {
    'Freelancer/Virtual Assistants': 'Cost: $120-$450/month | Problem: Inconsistent, unreliable, no continuity',
    'Generic Automation (Zapier, IFTTT)': 'Cost: $99-$300/month | Problem: Not specialized for healthcare, generic templates',
    'EHR Systems (OpenMRS, DHIS2)': 'Cost: $500-$5,000+/month | Problem: Expensive, complex, overkill for workflow automation',
    'In-House Development': 'Cost: $3,000-$10,000+ upfront + ongoing | Problem: Requires technical staff, maintenance burden',
}
for comp, details in competitors.items():
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(comp + ': ').bold = True
    p.add_run(details)

doc.add_heading('Our Competitive Advantage', 2)
advantages = [
    'Hyper-Specialization: Purpose-built for HIV care workflows + integrated comorbidity management (diabetes, hypertension, TB)',
    'Comprehensive Health Coverage: Chronic disease monitoring, maternal health tracking, medication adherence, appointment management - all in one platform',
    'Cost Leadership: 70% cheaper than EHR systems, more reliable than freelancers',
    'AI Integration: OpenAI-powered personalization, patient engagement, and executive summarization',
    'Flexibility: n8n-based architecture allows customization without vendor lock-in',
    'Patient-Centric: Improves adherence rates, reduces missed appointments, and improves clinical outcomes',
    'Proven Market Demand: Existing alternatives show clear willingness to pay ($120-$450/month) for similar single-purpose tools'
]
for adv in advantages:
    doc.add_paragraph(adv, style='List Bullet')

doc.add_heading('Potential Market Size (Target Region)', 2)
p = doc.add_paragraph(
    'Geographic Focus: Western Uganda (Mbarara, Fort Portal, Kabale, Kisoro, Kanungu, Rukungiri, Ntungamo Districts)'
)
justify_paragraph(p)
p = doc.add_paragraph(
    'Western Uganda has among the highest HIV prevalence rates in Uganda and Africa (12-18% in some districts, nearly 2x the national average). This concentrated high-burden region is our initial focus for rapid market penetration and relationship-building.'
)
justify_paragraph(p)
doc.add_paragraph()

market_tiers = {
    'Phase 1 - Western Uganda (Days 1-90)': {
        'Public Hospitals with HIV Programs': '3-4 facilities',
        'District Health Offices/Programs': '2-3 programs',
        'NGO Networks & Clinics': '2-3 networks',
        'Private Clinics': '1-2 facilities',
        'Academic/Teaching Institutions': '1 facility',
        'Phase 1 Target': '10-15 customers (conservative 50% conversion from 20 prospects)'
    },
    'Phase 2 - Western Uganda Expansion (Months 4-6)': {
        'Secondary District Facilities': '12-15 additional facilities',
        'Referral Network Growth': '10-15 additional prospects from existing customer referrals',
        'Phase 2 Target': '25-35 total customers in Western Uganda by Month 6'
    },
    'Long-term (Pan-Uganda Expansion)': {
        'Central Uganda (Kampala Metro)': '40-60 facilities',
        'Other Regions (after Western Uganda proof)': '100-150+ facilities',
        'Total Pan-Uganda TAM': '150-250+ HIV care facilities'
    }
}

for region, breakdown in market_tiers.items():
    doc.add_heading(region, 3)
    for key, value in breakdown.items():
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{key}: ').bold = True
        p.add_run(str(value))

doc.add_paragraph()
total_addressable = doc.add_paragraph()
total_addressable.add_run('Phase 1 Revenue Potential (Western Uganda, 90 days): ').bold = True
total_addressable.add_run('10-15 customers × UGX 2,956,300/month (avg Silver tier) = UGX 29,563,000-44,344,500 MRR by Day 90\n\n')
total_addressable.add_run('Phase 2 Revenue Potential (Extended Western Uganda, 6 months): ').bold = True
total_addressable.add_run('25-35 customers × UGX 2,956,300/month = UGX 73,907,500-103,470,500 MRR by Month 6\n\n')
total_addressable.add_run('Long-term Pan-Uganda TAM (12+ months): ').bold = True
total_addressable.add_run('150-250 facilities × UGX 2,956,300/month = UGX 443,445,000-739,075,000 MRR potential')

doc.add_page_break()

# SECTION 3: 30-60-90 DAY ACTION PLAN
doc.add_heading('3. 30-60-90 DAY ACTION PLAN', 1)

doc.add_heading('Days 1-30: Foundation & Initial Outreach', 2)

doc.add_heading('Deliverables', 3)
p = doc.add_paragraph()
foundation = [
    'Complete and test n8n workflow with real healthcare data (HIV + comorbidity management)',
    'Develop 14-day free trial setup process',
    'Create professional website with demo video',
    'Design branded presentation deck and case study template',
    'Identify and validate first 10 target customers (see Client Outreach Plan)',
    'Prepare personalized outreach messages for LinkedIn',
    'Build integrated modules: Chronic Disease Management, Maternal Health, Medication Adherence, Appointment Management'
]
for item in foundation:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Weekly Breakdown', 3)
doc.add_paragraph()
doc.add_paragraph('Week 1-2: Technical Setup').bold = True
doc.add_paragraph('Finalize n8n workflow with real-world sample data from HIV care facilities', style='List Bullet 2')
doc.add_paragraph('Build integrated modules: Chronic Disease tracking, Maternal Health monitoring, Medication Adherence system, Appointment Management', style='List Bullet 2')
doc.add_paragraph('Set up Google Sheets templates for all modules (HIV, CD, Maternal, Medications, Appointments)', style='List Bullet 2')
doc.add_paragraph('Configure OpenAI API and WhatsApp Business API integrations for patient engagement', style='List Bullet 2')
doc.add_paragraph('Complete n8n to Google Drive integration for automated report delivery across all modules', style='List Bullet 2')

doc.add_paragraph()
doc.add_paragraph('Week 2-3: Marketing & Content').bold = True
doc.add_paragraph('Design professional 1-page website highlighting key benefits', style='List Bullet 2')
doc.add_paragraph('Create 2-3 LinkedIn thought leadership posts on healthcare automation', style='List Bullet 2')
doc.add_paragraph('Record 3-5 minute demo video showing workflow in action', style='List Bullet 2')
doc.add_paragraph('Develop pricing page and comparison chart vs. competitors', style='List Bullet 2')

doc.add_paragraph()
doc.add_paragraph('Week 3-4: Prospect Research & Initial Outreach').bold = True
doc.add_paragraph('Research and identify 10 potential first customers (healthcare administrators/directors)', style='List Bullet 2')
doc.add_paragraph('Establish LinkedIn connections with decision-makers (Days 28-30)', style='List Bullet 2')
doc.add_paragraph('Prepare personalized outreach messages with clear value prop', style='List Bullet 2')
doc.add_paragraph('Schedule 5-8 discovery calls for Week 5', style='List Bullet 2')

doc.add_heading('Expected Outcome - Days 1-30', 3)
doc.add_paragraph('✓ 3-5 pilot customers in early talks (scheduled for demos Day 31+)')
doc.add_paragraph('✓ 50+ LinkedIn connections with qualified decision-makers')
doc.add_paragraph('✓ Fully functional workflow and trial setup')
doc.add_paragraph('✓ Website and marketing collateral live')

doc.add_page_break()

doc.add_heading('Days 31-60: Refinement, Early Growth, & First Pilots', 2)

doc.add_heading('Deliverables', 3)
growth = [
    'Conduct 5-8 discovery calls and personalized demos (HIV + comorbidity management)',
    'Gather feedback from pilot customers on all integrated modules (HIV, Chronic Disease, Maternal Health, Medication Adherence, Appointment Management)',
    'Secure 3-5 pilot customers with full platform access (discounted or free trial)',
    'Acquire 5-8 new qualified leads through content and outreach (emphasize integrated platform)',
    'Launch LinkedIn advertising campaign targeting healthcare administrators with emphasis on time-savings and clinical outcomes',
    'Develop first customer testimonial and case study showcasing module integration benefits',
    'Create detailed implementation guide for all integrated modules'
]
for item in growth:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Weekly Breakdown', 3)
doc.add_paragraph()
doc.add_paragraph('Week 5-6: Demo & Pilot Conversion').bold = True
doc.add_paragraph('Conduct 5-8 discovery calls to understand pain points and customize pitch', style='List Bullet 2')
doc.add_paragraph('Deliver personalized 30-minute demos with pilot data', style='List Bullet 2')
doc.add_paragraph('Secure commitment from 3-5 pilot customers (free/discounted trial)', style='List Bullet 2')
doc.add_paragraph('Send trial setup instructions and provide onboarding support', style='List Bullet 2')

doc.add_paragraph()
doc.add_paragraph('Week 7-8: Iteration & New Outreach').bold = True
doc.add_paragraph('Collect feedback from pilot customers (Week 5-6 trials now running)', style='List Bullet 2')
doc.add_paragraph('Iterate and optimize n8n workflows based on real feedback', style='List Bullet 2')
doc.add_paragraph('Launch 30-targeted LinkedIn ads with case study snippets', style='List Bullet 2')
doc.add_paragraph('Identify next 5-8 qualified leads from warm network and content engagement', style='List Bullet 2')
doc.add_paragraph('Schedule discovery calls with new prospects (Week 9+)', style='List Bullet 2')

doc.add_heading('Expected Outcome - Days 31-60', 3)
doc.add_paragraph('✓ 3-5 active pilot customers providing real-world feedback')
doc.add_paragraph('✓ 1-2 customers ready to convert to paid subscriptions (Silver tier)')
doc.add_paragraph('✓ 5-8 new qualified leads in demo pipeline')
doc.add_paragraph('✓ First customer testimonial and mini case study published')
doc.add_paragraph('✓ Optimized workflows based on pilot feedback')

doc.add_page_break()

doc.add_heading('Days 61-90: Scale & First Revenue Acceleration', 2)

doc.add_heading('Deliverables', 3)
scale = [
    'Convert 5-10 pilot and lead customers to paid subscriptions with integrated platform',
    'Achieve 15-20 total paying customers (break-even + growth)',
    'Publish first comprehensive customer case study with results across multiple modules',
    'Launch district-level and multi-facility outreach with module-specific value propositions',
    'Develop customer success playbook for onboarding all integrated modules',
    'Plan expansion to secondary products (advanced clinical analytics, patient portal) or regions',
    'Activate module upsells: Maternal Health, Chronic Disease Management, and Enhanced Appointment Management for existing customers'
]
for item in scale:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Weekly Breakdown', 3)
doc.add_paragraph()
doc.add_paragraph('Week 9-10: Conversion & Case Study').bold = True
doc.add_paragraph('Convert 3-5 active pilots to paid Silver tier customers', style='List Bullet 2')
doc.add_paragraph('Follow up with 5-8 demo pipeline prospects for decision', style='List Bullet 2')
doc.add_paragraph('Document 2-3 customer case studies with quantified results (e.g., "Saved 8 hours/week")', style='List Bullet 2')
doc.add_paragraph('Publish case study across website, LinkedIn, and email', style='List Bullet 2')

doc.add_paragraph()
doc.add_paragraph('Week 11-12: Expansion & Optimization').bold = True
doc.add_paragraph('Target district programs and multi-facility organizations', style='List Bullet 2')
doc.add_paragraph('Convert remaining demo leads (2-5 additional customers)', style='List Bullet 2')
doc.add_paragraph('Develop customer success documentation and self-serve onboarding guide', style='List Bullet 2')
doc.add_paragraph('Analyze MoM growth rate and plan next 90-day expansion', style='List Bullet 2')

doc.add_heading('Expected Outcome - Days 61-90', 3)
doc.add_paragraph('✓ 15-20 paying customers at UGX 2,956,300/month Silver tier')
doc.add_paragraph('✓ Monthly Recurring Revenue (MRR): UGX 44,344,500-UGX 59,126,000')
doc.add_paragraph('✓ Profitability achieved (break-even at 16 customers)')
doc.add_paragraph('✓ 2-3 published case studies demonstrating customer success')
doc.add_paragraph('✓ Clear growth momentum and expansion roadmap')
doc.add_paragraph('✓ 30-50+ leads in pipeline for next 90 days')

doc.add_page_break()

# SECTION 4: FINANCIAL PROJECTIONS
doc.add_heading('4. FINANCIAL PROJECTIONS', 1)

doc.add_heading('Pricing Structure', 2)
doc.add_paragraph()

pricing_table = doc.add_table(rows=5, cols=4)
pricing_table.style = 'Light Grid Accent 1'

# Header row
hdr_cells = pricing_table.rows[0].cells
hdr_cells[0].text = 'Tier'
hdr_cells[1].text = 'Monthly Price'
hdr_cells[2].text = 'Target Customer'
hdr_cells[3].text = 'Annual Value'

# Data rows (converted to UGX at 1 USD = 3,700 UGX)
tiers_data = [
    ('Bronze', 'UGX 1,476,300', 'Small clinics, NGOs', 'UGX 17,715,600'),
    ('Silver', 'UGX 2,956,300', 'Hospitals, district programs', 'UGX 35,475,600'),
    ('Gold', 'UGX 5,546,300', 'Multi-facility, regional programs', 'UGX 66,555,600'),
]
for i, (tier, price, customer, annual) in enumerate(tiers_data, start=1):
    row_cells = pricing_table.rows[i].cells
    row_cells[0].text = tier
    row_cells[1].text = price
    row_cells[2].text = customer
    row_cells[3].text = annual

doc.add_paragraph()
setup_fee = doc.add_paragraph()
setup_fee.add_run('One-Time Setup Fee: ').bold = True
setup_fee.add_run('UGX 7,400,000 per customer (covers initial workflow customization, data migration, training)')

doc.add_heading('Cost Breakdown - Silver Tier (UGX 2,956,300/month)', 2)

cost_table = doc.add_table(rows=7, cols=2)
cost_table.style = 'Light Grid Accent 1'

cost_rows = [
    ('Cost Category', 'Monthly Cost'),
    ('n8n Cloud Workspace (advanced)', 'UGX 185,000'),
    ('OpenAI API usage (1000+ messages/month)', 'UGX 92,500'),
    ('Google Workspace/Sheets (shared)', 'UGX 37,000'),
    ('WhatsApp Business API (per customer)', 'UGX 148,000'),
    ('Support & Onboarding (30 min/month)', 'UGX 555,000'),
    ('COGS Total', 'UGX 1,017,500'),
]

for i, (category, cost) in enumerate(cost_rows):
    cells = cost_table.rows[i].cells
    cells[0].text = category
    cells[1].text = cost

doc.add_paragraph()

margin_calc = doc.add_paragraph()
margin_calc.add_run('Gross Margin Calculation:\n').bold = True
margin_calc.add_run('Revenue: UGX 2,956,300/month\n')
margin_calc.add_run('COGS: UGX 1,017,500/month\n')
margin_calc.add_run('Gross Profit: UGX 1,938,800/month per customer\n')
margin_calc.add_run('Gross Margin: 65.6%')

doc.add_heading('Break-Even Analysis', 2)

breakeven = doc.add_paragraph()
breakeven.add_run('Fixed Costs (estimated monthly): UGX 18,500,000\n').bold = True
breakeven.add_run('  - Founder salary/living expenses: UGX 11,100,000\n')
breakeven.add_run('  - Marketing & content: UGX 4,440,000\n')
breakeven.add_run('  - Tools/infrastructure: UGX 2,960,000\n\n')
breakeven.add_run('Break-Even Point: \n').bold = True
breakeven.add_run('UGX 18,500,000 ÷ UGX 1,938,800 contribution per Silver customer = 9.5 customers\n')
breakeven.add_run('\nConservative Estimate: 12-16 Silver tier customers = profitability\n\n')
breakeven.add_run('Timeline: Achievable by Day 80-90 with 3-5 pilot conversions + 5-10 new sales')

doc.add_heading('Revenue Projection - First 6 Months', 2)

revenue_table = doc.add_table(rows=8, cols=4)
revenue_table.style = 'Light Grid Accent 1'

revenue_rows = [
    ('Month', 'Customers', 'MRR (Avg UGX 2,956,300)', 'Setup Fee Revenue'),
    ('Month 1 (Days 1-30)', '0', 'UGX 0', 'UGX 0'),
    ('Month 2 (Days 31-60)', '3-5', 'UGX 8,868,900-14,781,500', 'UGX 22,200,000-37,000,000'),
    ('Month 3 (Days 61-90)', '12-18', 'UGX 35,475,600-53,213,400', 'UGX 14,800,000-29,600,000'),
    ('Month 4', '18-25', 'UGX 53,213,400-73,907,500', 'UGX 14,800,000-29,600,000'),
    ('Month 5', '25-35', 'UGX 73,907,500-103,470,500', 'UGX 14,800,000-29,600,000'),
    ('Month 6', '35-45', 'UGX 103,470,500-133,033,500', 'UGX 7,400,000-22,200,000'),
]

for i, row_data in enumerate(revenue_rows):
    cells = revenue_table.rows[i].cells
    for j, text in enumerate(row_data):
        cells[j].text = text

doc.add_paragraph()
revenue_note = doc.add_paragraph()
revenue_note.add_run('Total Revenue Projection (6 months): ').bold = True
revenue_note.add_run('UGX 462,500,000 - UGX 740,000,000 (including setup fees)\n')
revenue_note.add_run('By Month 6 MRR: UGX 103,470,500 - UGX 133,033,500 (35-45 customers)')

doc.save(r'90-Day Action Plan - AI HIV Automation (Western Uganda).docx')
print('[OK] 90-Day Action Plan created successfully')
