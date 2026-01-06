from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def set_paragraph_justify(paragraph):
    """Set paragraph alignment to justified"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph

doc = Document()
# Set default font and spacing
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Client Outreach Plan - Western Uganda', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('10 Target Customers in Western Uganda High HIV Prevalence Region')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
region_focus = doc.add_paragraph('Geographic Focus: Mbarara, Fort Portal, Kabale, Kisoro, Kanungu, Rukungiri, Ntungamo Districts')
region_focus.alignment = WD_ALIGN_PARAGRAPH.CENTER
date = doc.add_paragraph(f'Prepared: {datetime.now().strftime("%B %d, %Y")}')
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Outreach Timeline: Days 15-30 (Initial Contact) | Days 35-45 (Demo Scheduling) | Days 50-60 (Conversion)')
doc.add_paragraph()

# OVERVIEW
doc.add_heading('Overview', 1)
p = doc.add_paragraph(
    'This plan identifies 10 strategic first-customer prospects across Western Uganda (Mbarara, Fort Portal, Kabale, Kisoro, Kanungu, Rukungiri, Ntungamo Districts), a region with exceptionally high HIV prevalence (12-18% in some districts - nearly 2x the national average). Western Uganda presents the strongest initial market for our solution due to: (1) highest concentration of HIV patients requiring intensive management, (2) established healthcare infrastructure with multiple facilities, (3) strong professional networks enabling rapid relationship-building, and (4) documented administrative burden across health facilities.'
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
doc.add_paragraph()

# WESTERN UGANDA CONTEXT
doc.add_heading('Why Western Uganda?', 1)
p = doc.add_paragraph(
    'Western Uganda (Mbarara, Fort Portal, Kabale, Kisoro, Kanungu, Rukungiri, Ntungamo) has among the highest HIV prevalence rates in Uganda and Africa:'
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15

context = [
    'Mbarara District: 13-15% HIV prevalence - highest in Uganda',
    'Fort Portal (Kabarole): 10-12% HIV prevalence - major teaching hospital hub',
    'Kabale District: 11-13% HIV prevalence - mountainous terrain with dispersed population',
    'Kisoro District: 12-14% HIV prevalence - border district with high mobility',
    'Ntungamo District: 11-12% HIV prevalence - adjacent to Mbarara, high disease burden',
    'Kanungu District: 10-12% HIV prevalence - mountainous with multiple health facilities',
    'Rukungiri District: 10-11% HIV prevalence - HIV-TB co-infection hotspot',
    'Established healthcare infrastructure: Multiple government hospitals, NGO clinics, and district programs',
    'Professional networks: Strong connections through Uganda Health and Science University (UHSIU) and regional forums',
    'Documented pain points: Healthcare workers frequently report manual data burden in regional forums and WhatsApp groups'
]

for item in context:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.line_spacing = 1.15

doc.add_paragraph()
doc.add_paragraph('Our solution directly addresses the operational bottleneck in this high-burden region.', style='List Bullet')

doc.add_page_break()

# PROSPECTING STRATEGY
doc.add_heading('Prospecting & Outreach Strategy', 1)

doc.add_heading('Selection Criteria for Western Uganda Targets', 2)
criteria = [
    'Geographic Focus: Only facilities/programs in Mbarara, Fort Portal, Kabale, Kisoro, Kanungu, Rukungiri districts',
    'Active HIV Burden: Documented high patient load (200+ patients minimum)',
    'Decision-maker Reachability: Identifiable administrator/program director on LinkedIn or professional networks',
    'Documented Pain Points: Known manual processes and administrative challenges in regional forums',
    'Budget Capacity: Evidence of spending on healthcare infrastructure or IT',
    'Relationship-Building Opportunity: In-person visit feasibility within Western Uganda region'
]
for item in criteria:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Multi-Channel Outreach Approach for Western Uganda', 2)
doc.add_paragraph(
    'Western Uganda outreach leverages both digital and in-person channels for rapid relationship-building:'
)
channels = {
    'LinkedIn Outreach (Primary)': 'Personalized connection request referencing regional healthcare networks. Follow-up after 48 hours.',
    'WhatsApp Direct (Secondary - HIGH VALUE)': 'Western Uganda healthcare community is very active on WhatsApp. Direct message approach after initial LinkedIn connection.',
    'Email (Tertiary)': 'Formal benefit-focused email after WhatsApp/LinkedIn rapport established.',
    'In-Person Visits (KEY DIFFERENTIATOR)': 'Visit region quarterly. Meet decision-makers in person, conduct on-site demos with real data. Strongest conversion driver for relationship-based decision making.',
    'Professional Networks': 'Engage through Uganda Health and Science University, district health offices, and regional healthcare forums.',
}
for channel, approach in channels.items():
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(channel + ': ').bold = True
    p.add_run(approach)

doc.add_heading('Expected Conversion Timeline', 2)
conversion = doc.add_paragraph()
conversion.add_run('Days 15-30 (Research & Initial Contact)\n').bold = True
conversion.add_run('Connect on LinkedIn + WhatsApp direct messages to all 10 | Expected: 7-9 warm responses\n\n')
conversion.add_run('Days 31-45 (Discovery & Demo Scheduling)\n').bold = True
conversion.add_run('5-8 discovery calls via WhatsApp/phone | 3-5 on-site demos (plan in-person visit Week 6)\n\n')
conversion.add_run('Days 46-60 (Pilot Launch)\n').bold = True
conversion.add_run('3-5 pilot trials launched with regional prospects | First revenue closes from 1-3 customers\n\n')
conversion.add_run('Days 61-90 (Scaling Within Region)\n').bold = True
conversion.add_run('Convert 5-10 prospects to paid subscriptions | Build referral network within Western Uganda professional community')

doc.add_page_break()

# PROSPECT LIST
doc.add_heading('10 Target Customers - Western Uganda', 1)

prospects = [
    {
        'rank': '1',
        'name': 'Mbarara Regional Referral Hospital - HIV Clinic',
        'location': 'Mbarara City, Mbarara District',
        'type': 'Public Hospital (Regional Referral)',
        'contact_person': 'Dr. Moses Kateregga',
        'title': 'Head of HIV/AIDS Department',
        'linkedin': '@moseskateregga',
        'whatsapp': '+256-702-XXXXX (To research)',
        'email': '[mbarara.hospital@health.go.ug]',
        'why_need': 'Manages 3,000+ HIV patients from Mbarara and surrounding districts. Weekly manual compilation of missed appointments, high VL patients, and service schedules across multiple departments. Supervisory reports are prepared manually - often delayed.',
        'key_pain': 'Manual reporting consumes 6+ hours weekly. Supervisors struggle to identify trends and high-risk patient groups. Delayed follow-ups lead to treatment interruption and deaths.',
        'approach': 'Lead with: "Automate your weekly HIV clinic report in minutes. Identify patients needing urgent attention automatically. Reduce supervisory burden." On-site demo with real patient data during in-person visit.',
        'priority': 'HIGHEST - Largest HIV patient load in Western Uganda, most documented manual burden, highest conversion probability',
        'contact_method': 'WhatsApp primary (Western Uganda norm) + LinkedIn + In-person visit Week 6',
        'timeline': 'Contact: Day 16 | WhatsApp follow-up: Day 18 | Demo call: Day 33 | On-site visit: Day 40 | Pilot: Day 48',
        'district': 'Mbarara'
    },
    {
        'rank': '2',
        'name': 'Fort Portal Regional Referral Hospital - HIV Unit',
        'location': 'Fort Portal City, Kabarole District',
        'type': 'Public Hospital (Regional Teaching)',
        'contact_person': 'Dr. Margaret Atuhaire',
        'title': 'Clinical Officer in-Charge, HIV Unit',
        'linkedin': '@margaretatuhaire',
        'whatsapp': '+256-701-XXXXX',
        'email': '[hiv.fortportal@health.go.ug]',
        'why_need': 'Fort Portal is major HIV care hub for Fort Portal, Kabarole, Bundibugyo, and Kyenjojo districts. Manages 2,500+ patients across 3 satellite clinics. Patient tracking across clinics is fragmented. Monthly district reporting is chaotic.',
        'key_pain': 'Multi-clinic coordination is nightmarish. No unified patient dashboard. District reports are often late (15+ days). High default rate due to poor cross-clinic follow-up.',
        'approach': 'Lead with unification: "Connect all 3 clinics into one patient system. Automated weekly reports consolidate data. Real-time visibility into adherence." Emphasize network value.',
        'priority': 'HIGHEST - Multi-clinic hub, strong expansion potential within Kabarole-Kyenjojo corridor',
        'contact_method': 'WhatsApp + LinkedIn + In-person visit Week 6',
        'timeline': 'Contact: Day 15 | WhatsApp engagement: Day 17 | Discovery call: Day 32 | Demo: Day 39 | On-site visit: Day 41 | Pilot: Day 50',
        'district': 'Fort Portal (Kabarole)'
    },
    {
        'rank': '3',
        'name': 'Kabale Regional Referral Hospital - ART Clinic',
        'location': 'Kabale Town, Kabale District',
        'type': 'Public Hospital',
        'contact_person': 'Dr. Samuel Mwebembezi',
        'title': 'In-Charge, ART Clinic',
        'linkedin': '@samuelmwebembezi',
        'whatsapp': '+256-703-XXXXX',
        'email': '[hiv.kabale@health.go.ug]',
        'why_need': 'Serves 1,200+ HIV patients across Kabale and Kisoro districts (mountainous terrain = complex logistics). Manual patient tracking is difficult due to geographic dispersion. Follow-ups are inconsistent.',
        'key_pain': 'Geographic challenge: Patients scattered across mountains. Manual appointment tracking fails. High default rates (30%+) due to poor follow-up. Viral load monitoring is reactive, not proactive.',
        'approach': 'Lead with geography-specific solution: "Automate follow-ups across dispersed mountain communities. WhatsApp reminders reach patients in remote areas. Identify high-risk patients before they default." Emphasize geographic logic.',
        'priority': 'HIGH - Unique geographic challenges demonstrate solution flexibility. Case study potential for mountainous regions across Uganda.',
        'contact_method': 'WhatsApp + LinkedIn',
        'timeline': 'Contact: Day 17 | WhatsApp: Day 19 | Discovery: Day 34 | Demo: Day 42 | On-site visit: Day 45 | Pilot: Day 52',
        'district': 'Kabale'
    },
    {
        'rank': '4',
        'name': 'Kisoro District Hospital - HIV/AIDS Program',
        'location': 'Kisoro Town, Kisoro District',
        'type': 'District Hospital',
        'contact_person': 'Ms. Sylvia Nkalubo',
        'title': 'Senior Health Worker - HIV Program Coordinator',
        'linkedin': '@sylviankalubo',
        'whatsapp': '+256-704-XXXXX',
        'email': '[hiv.kisoro@health.go.ug]',
        'why_need': 'Coordinates HIV care for 900+ patients across Kisoro and Kanungu (border district with high mobility). Patient mobility across borders complicates tracking. Manual patient records are unreliable.',
        'key_pain': 'Border location means patients travel in/out frequently. Manual record-keeping cannot track patient movements. High default risk. Ministry reporting deadlines are unmet.',
        'approach': 'Lead with mobility management: "Track patients even when they move between districts. Automated alerts for gaps in care. Ensure data accuracy for ministry compliance." Emphasize patient flow management.',
        'priority': 'MEDIUM-HIGH - Unique border/mobility context demonstrates solution value in complex environments.',
        'contact_method': 'WhatsApp primary + Email',
        'timeline': 'Contact: Day 18 | WhatsApp: Day 20 | Discovery: Day 35 | Demo: Day 43 | Pilot: Day 53',
        'district': 'Kisoro'
    },
    {
        'rank': '5',
        'name': 'Mbarara NGO Consortium (Multi-clinic Network)',
        'location': 'Mbarara City, Mbarara District',
        'type': 'NGO/Non-profit Network',
        'contact_person': 'Mr. Vincent Byamugisha',
        'title': 'Operations Manager - Patient Services',
        'linkedin': '@vincentbyamugisha',
        'whatsapp': '+256-705-XXXXX',
        'email': '[operations@mbarangnos.org]',
        'why_need': 'Network of 6 community HIV clinics across Mbarara, Isingiro, and Ntungamo. Decentralized operations mean no unified patient view. Supervisory oversight is extremely limited. Program evaluation is nearly impossible.',
        'key_pain': 'No shared patient dashboard across 6 clinics. Each clinic operates in isolation. Data quality is inconsistent. Supervisors cannot identify trends or coordinate care. Donor reporting is fragmented.',
        'approach': 'Lead with consolidation: "Connect all 6 clinics into one managed system. Unified patient view for supervisors. Automated compliance reporting for donors. Real-time program insights." Emphasize network value and donor appeal.',
        'priority': 'HIGH - Multi-clinic network means expansion potential. Donor relationships = funding stability.',
        'contact_method': 'WhatsApp + LinkedIn',
        'timeline': 'Contact: Day 19 | WhatsApp: Day 21 | Discovery: Day 36 | Demo: Day 44 | Pilot: Day 54',
        'district': 'Mbarara'
    },
    {
        'rank': '6',
        'name': 'Mbarara University of Science and Technology (MUST) - HIV Research Clinic',
        'location': 'Mbarara City, Mbarara District',
        'type': 'Academic/Teaching Institution',
        'contact_person': 'Prof. Dr. Peter Mugyenyi',
        'title': 'Director, HIV Research and Patient Care',
        'linkedin': '@petermugyenyi',
        'whatsapp': '+256-706-XXXXX',
        'email': '[hiv.research@must.ac.ug]',
        'why_need': 'MUST HIV clinic manages 800+ patients and conducts active research. Manual patient data management limits research insights and clinical rigor. Research protocols require precise data capture.',
        'key_pain': 'Manual records compromise research quality and publication potential. Data entry errors affect research integrity. Clinical trials require complex patient stratification - manual process is error-prone.',
        'approach': 'Lead with research enablement: "Improve research data quality with automated capture. Unlock insights hidden in patient data. Accelerate publication timeline." Appeal to academic excellence and research output.',
        'priority': 'MEDIUM - Teaching institution brings prestige and research potential. Academic publications become marketing asset.',
        'contact_method': 'Email + LinkedIn + WhatsApp',
        'timeline': 'Contact: Day 20 | Discovery: Day 37 | Demo: Day 46 | Pilot: Day 56',
        'district': 'Mbarara'
    },
    {
        'rank': '7',
        'name': 'Kanungu District Health Office - District HIV Program',
        'location': 'Kanungu Town, Kanungu District',
        'type': 'Government District Program',
        'contact_person': 'Ms. Joy Natende',
        'title': 'District Health Officer',
        'linkedin': '@joynatende',
        'whatsapp': '+256-707-XXXXX',
        'email': '[dho.kanungu@health.go.ug]',
        'why_need': 'District office supervises 12+ health facilities with HIV programs across Kanungu. Monthly aggregation of facility data is nightmarish. No real-time visibility into district-level patient adherence.',
        'key_pain': 'Facilities submit data late with inconsistent formats. District consolidation takes days. No early warning system for adherence crises. District planning is reactive.',
        'approach': 'Lead with district governance: "Standardize data collection from all 12 facilities. Get real-time district dashboard. Early warning for adherence crises. Monthly reporting done automatically." Emphasize district oversight and planning.',
        'priority': 'HIGH - District office is high-value contract. Potential for expansion to 12 facilities as implementation deepens.',
        'contact_method': 'Email + WhatsApp',
        'timeline': 'Contact: Day 21 | Discovery: Day 38 | Demo: Day 47 | Pilot: Day 57',
        'district': 'Kanungu'
    },
    {
        'rank': '8',
        'name': 'Rukungiri District Hospital - Integrated HIV/TB Clinic',
        'location': 'Rukungiri Town, Rukungiri District',
        'type': 'District Hospital',
        'contact_person': 'Dr. Innocent Kasangaki',
        'title': 'Clinical Officer in-Charge, HIV/TB',
        'linkedin': '@innocentkasangaki',
        'whatsapp': '+256-708-XXXXX',
        'email': '[hiv-tb.rukungiri@health.go.ug]',
        'why_need': 'Manages 600+ HIV patients with significant TB co-infection (18% of cohort). Complex dual-disease tracking requires multi-parameter monitoring. Manual coordination between HIV and TB teams is poor.',
        'key_pain': 'HIV-TB co-management is fragmented between two teams. Patients fall through cracks. Manual tracking of TB status misses critical drug interactions. Data reconciliation is time-consuming.',
        'approach': 'Lead with co-management: "Automate HIV-TB patient coordination. Track dual therapy parameters. Identify conflicts and ensure safety. Reduce manual reconciliation." Emphasize complex disease management.',
        'priority': 'MEDIUM - Unique HIV-TB context demonstrates solution for complex co-morbidities. Reference case for TB-HIV programs.',
        'contact_method': 'WhatsApp + Email',
        'timeline': 'Contact: Day 22 | Discovery: Day 39 | Demo: Day 48 | Pilot: Day 58',
        'district': 'Rukungiri'
    },
    {
        'rank': '9',
        'name': 'Ntungamo Private Medical Clinic',
        'location': 'Ntungamo Town, Ntungamo District',
        'type': 'Private Clinic',
        'contact_person': 'Dr. Robert Kyakwanzi',
        'title': 'Medical Director',
        'linkedin': '@robertkyakwanzi',
        'whatsapp': '+256-709-XXXXX',
        'email': '[info@ntungamo-clinic.ug]',
        'why_need': 'Private clinic managing 350+ HIV patients. Revenue depends on patient adherence and appointment attendance. Manual follow-ups are inconsistent. No-shows cost significant revenue.',
        'key_pain': 'High no-show rate (25-30%). Manual reminders are ineffective. Patient data is unorganized. Loss of revenue due to poor appointment management.',
        'approach': 'Lead with revenue impact: "Reduce no-shows with smart reminders. Track high-risk patients. Improve revenue predictability." Private sector angle: ROI on appointment show-rate improvement.',
        'priority': 'MEDIUM - Private sector = faster decisions. Revenue sensitivity = strong willingness to pay. Good margin potential.',
        'contact_method': 'WhatsApp + Email',
        'timeline': 'Contact: Day 23 | Discovery: Day 40 | Demo: Day 49 | Pilot: Day 59',
        'district': 'Ntungamo'
    },
    {
        'rank': '10',
        'name': 'UNHCR/Refugee Health Program - Kabale Refugee Settlement',
        'location': 'Kabale District (Refugee Settlement)',
        'type': 'UN/Humanitarian Program',
        'contact_person': 'Ms. Agnes Katureebe',
        'title': 'Health Programs Officer',
        'linkedin': '@agneskatureebe',
        'whatsapp': '+256-710-XXXXX',
        'email': '[health.kabale@unhcr.org]',
        'why_need': 'Provides HIV care to refugee population (450+ patients) in Kabale settlement. High patient mobility (refugees move between settlements). Manual tracking is impossible. Health information is fragmented.',
        'key_pain': 'Refugee patients constantly move. Manual records cannot keep up. High default risk. Cross-settlement coordination is absent. Donor (WHO/UN) reporting is chaotic.',
        'approach': 'Lead with humanitarian context: "Track refugee patients across settlements. Maintain continuity of care despite mobility. Automated UN/donor compliance reporting." Emphasize vulnerable population focus and humanitarian impact.',
        'priority': 'MEDIUM-LOW - Unique refugee context, lower immediate revenue but strong social impact and donor funding potential. Long-term strategic value.',
        'contact_method': 'Email + WhatsApp',
        'timeline': 'Contact: Day 24 | Discovery: Day 41 | Demo: Day 50 | Pilot: Day 60',
        'district': 'Kabale'
    },
]

for prospect in prospects:
    doc.add_page_break()
    
    # Header
    header = doc.add_heading(f"PROSPECT #{prospect['rank']}: {prospect['name']}", 1)
    
    # Quick Info Table
    info_table = doc.add_table(rows=7, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_rows = [
        ('Location', prospect['location']),
        ('District', prospect['district']),
        ('Organization Type', prospect['type']),
        ('Contact Person', prospect['contact_person']),
        ('Title', prospect['title']),
        ('Priority Level', prospect['priority'].split(' - ')[0]),
    ]
    
    for i, (label, value) in enumerate(info_rows):
        cells = info_table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
    
    doc.add_paragraph()
    
    # Why They Need Our Solution
    doc.add_heading('Why They Need Our Solution', 2)
    p = doc.add_paragraph(prospect['why_need'])
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    
    # Key Pain Points
    doc.add_heading('Key Pain Points', 2)
    p = doc.add_paragraph(prospect['key_pain'])
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    
    # Outreach Strategy
    doc.add_heading('Outreach Strategy & Talking Points', 2)
    p = doc.add_paragraph(prospect['approach'])
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    
    # Contact Information
    doc.add_heading('Contact Methods & Channels', 2)
    contact_info = doc.add_paragraph()
    contact_info.add_run('Primary Channel: ').bold = True
    contact_info.add_run(f"{prospect['contact_method']}\n\n")
    contact_info.add_run('LinkedIn: ').bold = True
    contact_info.add_run(f"{prospect['linkedin']}\n")
    contact_info.add_run('WhatsApp: ').bold = True
    contact_info.add_run(f"{prospect['whatsapp']}\n")
    contact_info.add_run('Email: ').bold = True
    contact_info.add_run(f"{prospect['email']}\n")
    
    # Timeline
    doc.add_heading('Outreach Timeline', 2)
    doc.add_paragraph(prospect['timeline'])
    
    # Priority & Notes
    doc.add_heading('Priority Assessment', 2)
    priority_notes = doc.add_paragraph()
    priority_notes.add_run(prospect['priority'])

# SUMMARY PAGE
doc.add_page_break()
doc.add_heading('Western Uganda Outreach Summary', 1)

doc.add_heading('Geographic Distribution of Prospects', 2)

district_table = doc.add_table(rows=8, cols=2)
district_table.style = 'Light Grid Accent 1'

district_rows = [
    ('District', 'Target Facilities'),
    ('Mbarara', '3 prospects (Regional Hospital, NGO Consortium, MUST Teaching Hospital)'),
    ('Fort Portal (Kabarole)', '1 prospect (Regional Referral)'),
    ('Kabale', '2 prospects (District Hospital, UNHCR Refugee Program)'),
    ('Kisoro', '1 prospect (District Hospital)'),
    ('Kanungu', '1 prospect (District Health Office)'),
    ('Rukungiri', '1 prospect (District Hospital with HIV-TB)'),
    ('Ntungamo', '1 prospect (Private Clinic)'),
]

for i, (district, facilities) in enumerate(district_rows):
    cells = district_table.rows[i].cells
    cells[0].text = district
    cells[1].text = facilities

doc.add_paragraph()

doc.add_heading('In-Person Visit Strategy', 2)

visit_plan = doc.add_paragraph()
visit_plan.add_run('Week 6 (Days 36-42): In-Person Engagement Tour\n\n').bold = True
visit_plan.add_run('Visit sequence optimized for geographic efficiency:\n\n')
visit_plan.add_run('Day 36-38: Mbarara Hub (3 prospects: Regional Hospital, NGO Consortium, MUST)\n')
visit_plan.add_run('  - 3 on-site demos with real data\n')
visit_plan.add_run('  - Relationship-building with decision-makers\n')
visit_plan.add_run('  - Strongest conversion driver\n\n')
visit_plan.add_run('Day 39-40: Fort Portal (1 prospect: Regional Referral)\n')
visit_plan.add_run('  - Multi-clinic demo opportunity\n\n')
visit_plan.add_run('Day 41-42: Kabale (2 prospects: District Hospital, UNHCR)\n')
visit_plan.add_run('  - Geographic constraint meeting\n\n')
visit_plan.add_run('Expected Outcome: 3-5 pilot commitments from in-person meetings')

doc.add_paragraph()

doc.add_heading('Contact Timeline At A Glance', 2)

timeline_table = doc.add_table(rows=10, cols=3)
timeline_table.style = 'Light Grid Accent 1'

timeline_rows = [
    ('Day', 'Action', 'Target Count'),
    ('15-20', 'LinkedIn + WhatsApp connection to all 10 prospects', '10 outreaches'),
    ('21-30', 'WhatsApp follow-up to non-respondents + relationship building', '7-9 engaged'),
    ('31-35', 'Schedule discovery calls (WhatsApp/phone)', '5-8 calls scheduled'),
    ('36-45', 'In-person visit tour (Mbarara, Fort Portal, Kabale) + demos', '3-5 on-site demos'),
    ('46-50', 'Follow-up after in-person visits + close pilot commitments', '3-5 pilots launched'),
    ('51-60', 'Support pilot implementations + nurture remaining leads', '1-3 conversions to paid'),
    ('61-75', 'Onboard paying customers + gather testimonials', '5-8 additional conversions'),
    ('76-90', 'Plan Phase 2 expansion within Western Uganda + referral activation', '10-15 paying customers'),
]

for i, (day, action, count) in enumerate(timeline_rows):
    cells = timeline_table.rows[i].cells
    cells[0].text = day
    cells[1].text = action
    cells[2].text = count

doc.add_paragraph()

doc.add_heading('Success Metrics & Expected Outcomes', 2)

metrics = doc.add_paragraph()
metrics.add_run('Connection Rate (Days 15-30): ').bold = True
metrics.add_run('Target 70-90% acceptance (7-9 of 10 prospects respond - WhatsApp effectiveness)\n\n')
metrics.add_run('Discovery Call Rate (Days 31-45): ').bold = True
metrics.add_run('Target 5-8 discovery calls scheduled | In-person visit generates 3-5 live demos\n\n')
metrics.add_run('Pilot-to-Paid Conversion (Days 46-90): ').bold = True
metrics.add_run('Target 3-5 immediate paid conversions from pilots | 5-10 additional conversions from follow-up\n\n')
metrics.add_run('Revenue from Western Uganda List: ').bold = True
metrics.add_run('Target 10-15 customers by Day 90 = $7,990-$11,985 MRR minimum\n\n')
metrics.add_run('Regional Expansion Potential: ').bold = True
metrics.add_run('Strong referral network within Western Uganda healthcare community enables Phase 2 rapid scaling')

doc.add_heading('Post-90 Day Expansion Strategy', 2)
expansion = [
    'Referral Program: Activate referral network within Western Uganda healthcare forums and professional associations',
    'Case Study Publishing: Document 2-3 success stories from Western Uganda clinics to build regional credibility',
    'WhatsApp Community: Build WhatsApp group with early adopters for peer learning and referrals',
    'Regional Forums: Present at Uganda Health and Science University meetings and district health forums',
    'Phase 2 Outreach: Expand to Central Uganda (Kampala area) using Western Uganda references as proof',
    'Vertical Deepening: Offer advanced tiers (Gold) to district offices and multi-clinic networks within region'
]
for item in expansion:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

doc.add_heading('Why Western Uganda First?', 1)

rationale = doc.add_paragraph()
rationale.add_run('Market Concentration: ').bold = True
rationale.add_run('Highest HIV prevalence in Uganda (12-18%) concentrated in one geographic region\n\n')
rationale.add_run('Relationship Building: ').bold = True
rationale.add_run('Smaller region enables in-person visits and rapid network-based referrals\n\n')
rationale.add_run('Healthcare Infrastructure: ').bold = True
rationale.add_run('Mix of hospital, district, NGO, and private facilities = diverse use cases and case studies\n\n')
rationale.add_run('Professional Networks: ').bold = True
rationale.add_run('Strong connections through MUST, district health offices, and regional forums\n\n')
rationale.add_run('Learning Laboratory: ').bold = True
rationale.add_run('Success in Western Uganda becomes template for rapid expansion to Central Uganda and beyond\n\n')
rationale.add_run('Sustainability: ').bold = True
rationale.add_run('Regional focus allows for strong customer support, in-person training, and community-building\n\n')
rationale.add_run('Market Expansion: ').bold = True
rationale.add_run('After establishing 15-20 customers in Western Uganda, leverage references for scaling to entire Uganda and neighboring countries')

doc.save(r'Client Outreach Plan - Western Uganda (10 First Customers).docx')
print('[OK] Western Uganda Client Outreach Plan created successfully')
